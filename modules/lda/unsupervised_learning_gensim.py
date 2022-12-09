# -*- coding: utf-8 -*-
"""
Created on Sun May  3 00:02:29 2020

@author: cvicentm
"""

# -*- coding: utf-8 -*-
"""
Created on Fri Apr 17 07:46:39 2020

@author: cvicentm
"""


#import pyLDAvis.sklearn
import pickle
import timeit
import datetime
import gensim, spacy, logging, warnings
import gensim.corpora as corpora
from gensim.utils import lemmatize, simple_preprocess
from gensim.models import CoherenceModel
from PIL import ImageColor
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from nltk import word_tokenize
from random import randint
import os
from tqdm import tqdm
import matplotlib.pyplot as plt
import numpy as np
from datetime import datetime
import ast
from itertools import chain
from scipy.signal import savgol_filter

#started importation my own modules---------------------------------
from modules.pre.create_corpus import create_corpus
from modules.pre.create_corpus import normalize_word
from modules.sql import dBAdapter

#finised importation my own modules---------------------------------



name_log_file = datetime.now().strftime('logs\key_error_%d_%m_%Y.log')
    
logging.basicConfig(filename=name_log_file, level=logging.WARNING, 
                    format="%(asctime)s:%(filename)s:%(lineno)d:%(levelname)s:%(message)s")



def printColorWordDocument(number,colors,generator_normalize,dic_subtitles,lda_model,corpus,n_documents):
    #make an explanation of every of the parameters of this function
    corp_cur = corpus[number] 
    topic_percs, wordid_topics, wordid_phivalues = lda_model[corp_cur]
    word_dominanttopic = [(lda_model.id2word[wd], topic[0]) for wd, topic in wordid_topics if topic]
                          
    
    def wordToDocument (word, color_hex):
        color_rgb = ImageColor.getrgb(color_hex)
        run=paragraph.add_run(word+" ")
        font = run.font
        font.color.rgb = RGBColor(color_rgb[0], color_rgb[1], color_rgb[2])
        """
        if color_hex != colors[len(colors)-1]:
            font.highlight_color = WD_COLOR_INDEX.YELLOW
        else:
            font.highlight_color = WD_COLOR_INDEX.WHITE
        """
    document_classified=[]
    word_dominanttopic_dict=dict(word_dominanttopic)
    dict_one_subtitle_token=word_tokenize(dic_subtitles[list(dic_subtitles.keys())[number]])
    for word in dict_one_subtitle_token:
        topic_word=len(colors)-1
        try:
            if normalize_word(word) in generator_normalize[number]:
                try:
                    topic_word=word_dominanttopic_dict[normalize_word(word)]
                except KeyError as error:
                    # Escribir aquí un log en vez de un print
                    logging.warning("OSError --- key error: "+str(word))
                    
            else:
                topic_word=len(colors)-1
            document_classified.append((word,topic_word))
        except:
            print("errrror")
    
    
    #document_classified=[(word,word_dominanttopic_dict[word]) for word in generator_normalize[0]]
    document = Document()
    paragraph = document.add_paragraph()
    [wordToDocument(word,colors[topic]) for word,topic in document_classified]
    
    word_subtitles_colors="word\\"+str(n_documents)+"\\"+list(dic_subtitles.keys())[number]+".docx"
    document.save(word_subtitles_colors)
    

def training_model(n_documents,n_topics,id2word, corpus, generator_normalize):
    #ESCRIBIR QUE HACE ESTA FUNCIÓN Y PARA QUE SIRVE CADA UNO DE SUS PARÁMETROS
    
    
    print("the model is being trained with: "+str(n_topics)+ "topics")
    
    lda_model = gensim.models.ldamodel.LdaModel(corpus=corpus,
                                               id2word=id2word,
                                               num_topics=n_topics, 
                                               random_state=100,
                                               update_every=1,
                                               chunksize=100,
                                               passes=10,
                                               alpha='auto',
                                               per_word_topics=True)
    
    
    #tengo bastantes dudas si en el parámetro texts tengo que poner generator_normalize o no
    coherencemodel = CoherenceModel(model=lda_model, corpus=corpus, dictionary=id2word, coherence='u_mass')
    #coherencemodel_cv = CoherenceModel(model=lda_model, texts=list(generator_normalize), dictionary=id2word, coherence='c_v')
    coherence_values = coherencemodel.get_coherence()
    #coherencemodel_c_uci = CoherenceModel(model=lda_model, texts=list(generator_normalize), dictionary=id2word, coherence='c_uci')

    #the model is gonna be saved
    # if the number of subtitles doest change, we can use the same model than the last time
    file_lda_model = 'D:\\caleb\\pickle\\'+str(n_documents)+'\\coherence_lda\\lda_model_'+str(n_topics)+'_'+str(n_documents)+'.sav'
    pickle.dump(lda_model, open(file_lda_model, 'wb'))
    
    return coherence_values


#PROGRAM......................................................................
def LDAmodel( n_topics, n_documents, n_printedDocuments,name_database,name_collection, step=1, start=1):   
    #Tengo que escribir para que sirve cada cosa que hace el gensim
    
    #import from DDBB dic_subtitles and generator normalize--------------------
    """
    print("Getting body subtitles from the database started ...")
    dbAdapter= dBAdapter.Database()
    dbAdapter.open()
    dic_subtitles = dict(dbAdapter.selectDic_subtitles_limit(n_documents))
    gn = dbAdapter.selectGenerator_normalize_limit(n_documents)
    generator_normalize = [ast.literal_eval(gni[0]) for gni in gn]
    dbAdapter.close()
    print("Getting body subtitles from the database finished ...")
    """
    print("Getting body subtitles from the database started ...")
    dbAdapter= dBAdapter.Database(name_database,name_collection)
    dbAdapter.open()
    listado=dbAdapter.selectGenerator_normalize_limit(n_documents)
    dic_subtitles = dbAdapter.selectDic_subtitles_limit(n_documents)
    dbAdapter.close()
    print("finalizada consulta")
    
    dic_subtitles2 = dic_subtitles
    
    generator_normalize = []
    for i in range(len(listado)):
        try:
            generator_normalize.append(listado[i].split(","))
        except:
            dic_subtitles2.pop(list(dic_subtitles.keys())[i])
            print("generator NonType------>"+str(i))
    
    dic_subtitles = dic_subtitles2
            
    for gn in generator_normalize:
        while True:
            try:
                gn.remove("")
            except ValueError:
                break
    print("Getting body subtitles from the database finished ...")
    n_documents = len(generator_normalize)
    #--------------------------------------------------------------------------
    coherencemodelArray = []
    perplexitymodelArray = []
    boundArray = []
    
    if not os.path.exists('D:\\caleb\\pickle\\'+str(n_documents)):
        os.makedirs('D:\\caleb\\pickle\\'+str(n_documents))
    try: 
        
        id2word = pickle.load(open("D:\\caleb\\pickle\\"+str(n_documents)+"\\coherence_lda\\id2word_"+str(n_documents)+".txt", "rb"))
        corpus = pickle.load(open("D:\\caleb\\pickle\\"+str(n_documents)+"\\coherence_lda\\corpus_"+str(n_documents)+".txt", "rb"))
        print("generator_normalize, id2word and corpus has been imported")
    
    except IOError:
        
        print("Proccess of creating corpus and the dictionary has started")
        #this is creating a dictionary with all de different words of the document
        id2word = corpora.Dictionary(generator_normalize)
        file_id2word = "D:\\caleb\\pickle\\"+str(n_documents)+"\\coherence_lda\\id2word_"+str(n_documents)+'.txt'
        pickle.dump(id2word, open(file_id2word, 'wb'))
        # Create Corpus: Term Document Frequency
        corpus = [id2word.doc2bow(text) for text in generator_normalize]
        file_corpus = "D:\\caleb\\pickle\\"+str(n_documents)+"\\coherence_lda\\corpus_"+str(n_documents)+'.txt'
        pickle.dump(corpus, open(file_corpus, 'wb'))
        
        print("Proccess of creating corpus and the dictionary has ended")
    
    for n_topics in chain(range(1,21,1),range(27,200,8)):
        #for n_topics in range(25,26):
        file_lda_model = 'D:\\caleb\\pickle\\'+str(n_documents)+'\\coherence_lda\\lda_model_'+str(n_topics)+'_'+str(n_documents)+'.sav'
        try:
           
            f=open(file_lda_model, 'rb')
            lda = pickle.load(f)
            print("The model has been trained previously with..."+str(n_topics)+" n_topics") 
            coherencemodel = CoherenceModel(model=lda, corpus=corpus, dictionary=id2word, coherence='u_mass')
            #coherencemodel_cv = CoherenceModel(model=lda, texts=list(generator_normalize), dictionary=id2word, coherence='c_v')
            #coherencemodel_c_uci = CoherenceModel(model=lda, texts=list(generator_normalize), dictionary=id2word, coherence='c_uci')
            #file_coherence_cv = 'D:\\caleb\\pickle\\'+str(n_documents)+'\\coherence_lda\\cv_'+str(n_topics)+'_'+str(n_documents)+'.sav'
            #pickle.dump(coherencemodel_cv, open(file_coherence_cv, 'wb'))
            #file_coherence_c_uci = 'D:\\caleb\\pickle\\'+str(n_documents)+'\\coherence_lda\\c_uci_'+str(n_topics)+'_'+str(n_documents)+'.sav'
            #pickle.dump(coherencemodel_c_uci, open(file_coherence_c_uci, 'wb'))
            #CoherenceModel(model=goodLdaModel, texts=texts, dictionary=dictionary, coherence='c_v')
            #coherencemodel = CoherenceModel(model=lda, texts=list(generator_normalize), dictionary=id2word, coherence='c_v')
            coherence_values = coherencemodel.get_coherence()
            coherencemodelArray.append(coherence_values)
            
            perplexitymodelArray.append(lda.log_perplexity(corpus))
            boundArray.append(lda.bound(corpus))
            
        except IOError:
            
            print("FINALLY: the LDA model has to be trained for "+str(n_documents)+" n_documents and "+str(n_topics)+" n_topics, trained")
            
            tic_all_processing=timeit.default_timer()
            #function based on : https://www.machinelearningplus.com/nlp/topic-modeling-gensim-python/#13viewthetopicsinldamodel
            [coherence_values]=training_model(n_documents,n_topics,id2word,corpus,generator_normalize)
            coherencemodelArray.append(coherence_values)  
            toc_all_processing=timeit.default_timer()
            try: 
                time_lda_fit=str(datetime.timedelta(seconds=int(float(toc_all_processing-tic_all_processing))))
                print("The process of training lda model with "+str(n_topics)+" n_topics and "+str(n_documents)+" n_documents, has taken "+time_lda_fit+" seconds")    
            except AttributeError: 
                print("The process of training lda model with "+str(n_topics)+" n_topics and "+str(n_documents)+" n_documents, has ended")
            
            #file_coherence_cv = 'D:\\caleb\\pickle\\'+str(n_documents)+'\\coherence_lda\\cv_'+str(n_topics)+'_'+str(n_documents)+'.sav'
            #pickle.dump(coherencemodel_cv, open(file_coherence_cv, 'wb'))
            #file_coherence_c_uci = 'D:\\caleb\\pickle\\'+str(n_documents)+'\\coherence_lda\\c_uci_'+str(n_topics)+'_'+str(n_documents)+'.sav'
            #pickle.dump(coherencemodel_c_uci, open(file_coherence_c_uci, 'wb'))
    
    
    coherencemodelArray = list(coherencemodelArray)
    file_coherence_umass = 'D:\\caleb\\pickle\\coherencemodelarray.sav'
    pickle.dump(coherencemodelArray, open(file_coherence_umass, 'wb'))
    x = list(chain(range(1,21,1),range(27,200,8)))
    #n_topics+1 because has to have the same weight than coherencemodelArray
    #score = savgol_filter(coherencemodelArray, 11, 3)
    plt.plot(x, coherencemodelArray)
    plt.xlabel("N_Topics")
    plt.ylabel("Coherence")
    plt.legend(("coherence_values"), loc='best')
    plt.show()
    
    best_n_topic=coherencemodelArray.index(min(coherencemodelArray))+start
    best_n_topic = best_n_topic * 2
    best_n_topic = 27
    
    
    print("el mejor modelo es: "+'pickle'+str(n_documents)+'\lda_model_'+str(best_n_topic)+'_'+str(n_documents)+'.sav')
    f=open('D:\\caleb\\pickle\\'+str(n_documents)+'\\coherence_lda\\lda_model_'+str(best_n_topic)+'_'+str(n_documents)+'.sav', 'rb')
    lda = pickle.load(f)
    document_per_topic=list(lda.get_document_topics(corpus))
    """
    corp_cur = corpus[1]
    topic_percs, wordid_topics, wordid_phivalues = lda[corp_cur]
    print(wordid_topics)
    """
    array_topic_per_document = np.zeros((len(document_per_topic), best_n_topic))
    
    
    for i in range(len(document_per_topic)):
        for j in range(len(document_per_topic[i])):
            try:    
                array_topic_per_document[i][document_per_topic[i][j][0]]= document_per_topic[i][j][1]
            except IndexError as index:
                #EN ESTE LOG sería necesario ponerle, cual ha sido el subtítulo que ha dado problemas e identifcar porque
                logging.warning("array_topic_per_document out of range in position n_document: "+str(i)+" and topic: "+str(j)+" \n")
    #NUMBER OF DOCUMENTs to print results on word
    
    
    return array_topic_per_document, best_n_topic, dic_subtitles,lda,generator_normalize,corpus,id2word,coherencemodelArray, perplexitymodelArray,boundArray


